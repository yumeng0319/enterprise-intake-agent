"""Build a compact 1-page design_doc.docx from the case study content."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x4F, 0x46, 0xE5)
MUTED = RGBColor(0x55, 0x55, 0x55)
TEXT = RGBColor(0x1A, 0x1A, 0x1A)

doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.45)
section.bottom_margin = Inches(0.45)
section.left_margin = Inches(0.55)
section.right_margin = Inches(0.55)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(9)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.05


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_title():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("Enterprise Intake Agent — System Design")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = TEXT

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    run2 = p2.add_run("Enterprise AI Builder Case Study  ·  One-page system design document")
    run2.italic = True
    run2.font.size = Pt(8.5)
    run2.font.color.rgb = MUTED


def add_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = ACCENT


def add_body(text, space_after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    _add_runs(p, text)
    return p


def _add_runs(p, text):
    """Very small **bold** markdown parser so body text can bold key terms/numbers."""
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.size = Pt(9)
        run.font.color.rgb = TEXT
        if i % 2 == 1:
            run.bold = True


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.18)
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        _add_runs(p, item)


def add_flow_diagram():
    stages = ["1 Extract\n(LLM)", "2 Classify\n(LLM)", "3 Priority\n(rule engine)", "4 Route\n(LLM+policy)", "5 Respond\n(LLM)"]
    table = doc.add_table(rows=1, cols=len(stages))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    col_width = Inches(1.34)
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width
    widths_dxa = int(col_width.twips)
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    for gridCol in tblGrid.findall(qn("w:gridCol")):
        gridCol.set(qn("w:w"), str(widths_dxa))

    for i, stage in enumerate(stages):
        cell = table.rows[0].cells[i]
        cell.width = col_width
        set_cell_shading(cell, "EEF2FF" if i != 2 else "FEF3C7")
        cell.vertical_alignment = 1
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        lines = stage.split("\n")
        run = p.add_run(lines[0])
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = ACCENT if i != 2 else RGBColor(0x92, 0x40, 0x0E)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        run2 = p2.add_run(lines[1])
        run2.font.size = Pt(7.5)
        run2.font.color.rgb = MUTED
    return table


# ---------------------------------------------------------------------------

add_title()

add_heading("Architecture & Reasoning")
add_body(
    "Five sequential stages, each stage's output feeding the next. Core principle: "
    "**the LLM handles language understanding; the highest-stakes decision (priority) "
    "is handled by deterministic code, not model judgment** — so “why P1?” always "
    "has an auditable answer, never “the model thought so.”",
    space_after=4,
)
add_flow_diagram()
add_body(
    "A policy layer forces mandatory human review whenever: domain is Security, priority "
    "is P0, the requester makes an unverified seniority claim (“I'm the CFO”), or "
    "completeness_score < 0.5. Unverified identity claims can raise priority out of the P3 "
    "default only alongside real stated business context — never on a bare claim — and "
    "never bypass human review either way.",
    space_after=2,
)

add_heading("Input / Reasoning Steps / Output")
add_body(
    "**Input:** one raw, unstructured text request (Slack/email-style, no form fields). "
    "**Steps:** extract facts (system_down, deadline_urgency, scope, business_impact, "
    "claimed_seniority, completeness_score) → classify domain/work type → compute "
    "priority from stage-1 facts via rule engine → recommend a team from a domain→team "
    "map → draft the first reply, or a clarification-only reply if completeness is too "
    "low to route confidently. **Output:** classification, priority (P0–P3 + rule fired "
    "+ reasoning), routing recommendation (team + human-review flag + reason), draft "
    "response.",
)

add_heading("Validation")
add_body(
    "16-case labeled eval set spanning every domain/work type, expected priority computed "
    "by running hand-specified facts through the real rule engine. **Domain 16/16 (100%), "
    "work type 16/16 (100%), priority 14/16 (88%)** — both misses traced to imprecise "
    "eval labels, not model error. needs_human_review raw agreement was 7/16, but "
    "**9 of 9 disagreements were the system being more cautious than the naive label "
    "expected, zero were under-cautious** — not a rigorous benchmark, but real evidence "
    "the system's error mode skews toward caution, not false confidence.",
)

add_heading("What's Automated vs. Left to Humans — and Why")
add_body(
    "**Automated:** understanding raw text, domain/work-type classification, team routing, "
    "response drafting, and the priority computation itself. **Left to humans:** final "
    "action on P0-Critical, Security-domain, low-completeness, or unverified-authority "
    "requests — precisely where an autonomous wrong call is expensive or the system "
    "lacks enough signal to act. Knowing when not to decide is as important as deciding "
    "well.",
)

add_heading("Assumptions Made — and What Breaks at Scale")
add_bullets([
    "**Text-only input** — real intake includes screenshots/attachments, not handled.",
    "**Routing table** is reverse-engineered from public job postings, not a confirmed org chart.",
    "**Priority thresholds are hand-picked**, not calibrated against real historical triage data.",
    "**No memory across requests** — recurring issues look like unrelated one-offs.",
    "**Sequential LLM calls add ~20–30s/request latency** — fine for a demo, not high-volume real-time.",
    "**LLM structured-output reliability isn't 100%** — one schema shape produced malformed tool output ~50% of calls; fixed by restructuring the schema + required-field validation/retry.",
])

add_heading("V2 — and What We'd Need to Know Before Building It")
add_body(
    "**V2 adds:** persistent state for recurring/duplicate issues, real identity "
    "verification (SSO/directory) replacing self-reported claims, multi-turn "
    "clarification, async processing. A toy version of one already works: logged human "
    "corrections trigger an advisory note when a rule is repeatedly overridden; real V2 "
    "would log every firing to compute an actual override rate and auto-adjust "
    "thresholds. **Before building it:** real historical triage data, access to an "
    "identity/directory system, and actual volume/latency requirements from the owning "
    "team.",
    space_after=0,
)

doc.save("/Users/yumengwang/Desktop/Snowflake AI Product/intake_agent/design_doc.docx")
print("Saved design_doc.docx")
